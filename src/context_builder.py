from __future__ import annotations

import csv
import json
import re
from collections.abc import Iterable
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".csv",
    ".tsv",
    ".xlsx",
    ".xls",
    ".json",
    ".jsonl",
}


@dataclass
class ContextChunk:
    id: str
    source_path: str
    source_name: str
    source_type: str
    section: str
    chunk_index: int
    text: str
    metadata: dict[str, Any]


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size.")

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        candidate = text[start:end]

        if end < len(text):
            split_at = max(candidate.rfind("\n\n"), candidate.rfind(". "), candidate.rfind("; "))
            if split_at > chunk_size * 0.5:
                end = start + split_at + 1
                candidate = text[start:end]

        chunks.append(candidate.strip())
        if end >= len(text):
            break
        start = max(0, end - chunk_overlap)

    return [chunk for chunk in chunks if chunk]


def extract_pdf(path: Path) -> Iterable[tuple[str, str, dict[str, Any]]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        yield f"page {index}", text, {"page": index}


def extract_docx(path: Path) -> Iterable[tuple[str, str, dict[str, Any]]]:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_blocks: list[str] = []

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            table_blocks.append(f"Table {table_index}\n" + "\n".join(rows))

    yield "document", "\n\n".join(paragraphs + table_blocks), {}


def extract_text_file(path: Path) -> Iterable[tuple[str, str, dict[str, Any]]]:
    yield "document", path.read_text(encoding="utf-8", errors="replace"), {}


def dataframe_to_text(df: Any, max_rows: int) -> str:
    if len(df) > max_rows:
        df = df.head(max_rows)
    df = df.fillna("")
    return df.to_csv(index=False, quoting=csv.QUOTE_MINIMAL)


def extract_csv(path: Path, max_rows: int) -> Iterable[tuple[str, str, dict[str, Any]]]:
    import pandas as pd

    separator = "\t" if path.suffix.lower() == ".tsv" else ","
    df = pd.read_csv(path, sep=separator)
    yield "table", dataframe_to_text(df, max_rows), {"rows": min(len(df), max_rows), "columns": list(df.columns)}


def extract_excel(path: Path, max_rows: int) -> Iterable[tuple[str, str, dict[str, Any]]]:
    import pandas as pd

    workbook = pd.ExcelFile(path)
    for sheet_name in workbook.sheet_names:
        df = workbook.parse(sheet_name)
        yield (
            f"sheet {sheet_name}",
            dataframe_to_text(df, max_rows),
            {"sheet": sheet_name, "rows": min(len(df), max_rows), "columns": list(df.columns)},
        )


def extract_json(path: Path) -> Iterable[tuple[str, str, dict[str, Any]]]:
    if path.suffix.lower() == ".jsonl":
        lines = []
        for line_number, line in enumerate(path.read_text(encoding="utf-8", errors="replace").splitlines(), start=1):
            if line.strip():
                lines.append(json.dumps(json.loads(line), ensure_ascii=False))
            if line_number % 500 == 0:
                yield f"lines up to {line_number}", "\n".join(lines), {"line_end": line_number}
                lines = []
        if lines:
            yield "remaining lines", "\n".join(lines), {}
        return

    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    yield "document", json.dumps(data, ensure_ascii=False, indent=2), {}


def extract_file(path: Path, max_rows_per_sheet: int) -> Iterable[tuple[str, str, dict[str, Any]]]:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".docx":
        return extract_docx(path)
    if extension in {".txt", ".md"}:
        return extract_text_file(path)
    if extension in {".csv", ".tsv"}:
        return extract_csv(path, max_rows_per_sheet)
    if extension in {".xlsx", ".xls"}:
        return extract_excel(path, max_rows_per_sheet)
    if extension in {".json", ".jsonl"}:
        return extract_json(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def iter_source_files(source_dir: Path, include_extensions: set[str]) -> Iterable[Path]:
    for path in sorted(source_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in include_extensions:
            yield path


def build_context_chunks(
    source_dir: Path,
    include_extensions: set[str] | None = None,
    chunk_size: int = 1800,
    chunk_overlap: int = 250,
    max_rows_per_sheet: int = 5000,
) -> list[ContextChunk]:
    include_extensions = include_extensions or SUPPORTED_EXTENSIONS
    chunks: list[ContextChunk] = []

    for path in iter_source_files(source_dir, include_extensions):
        source_type = path.suffix.lower().lstrip(".")
        relative_path = path.relative_to(source_dir).as_posix()

        for section, text, metadata in extract_file(path, max_rows_per_sheet):
            for chunk_index, chunk in enumerate(chunk_text(text, chunk_size, chunk_overlap), start=1):
                chunk_id = f"{relative_path}::{section}::{chunk_index}"
                chunks.append(
                    ContextChunk(
                        id=chunk_id,
                        source_path=relative_path,
                        source_name=path.name,
                        source_type=source_type,
                        section=section,
                        chunk_index=chunk_index,
                        text=chunk,
                        metadata=metadata,
                    )
                )

    return chunks


def write_jsonl(chunks: list[ContextChunk], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as file:
        for chunk in chunks:
            file.write(json.dumps(asdict(chunk), ensure_ascii=False) + "\n")


def write_markdown_pack(
    chunks: list[ContextChunk],
    output_path: Path,
    title: str,
    default_instruction: str,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    lines = [f"# {title}", "", default_instruction, ""]

    for chunk in chunks:
        lines.extend(
            [
                f"## {chunk.source_path} - {chunk.section} - chunk {chunk.chunk_index}",
                "",
                f"- id: `{chunk.id}`",
                f"- type: `{chunk.source_type}`",
                "",
                chunk.text,
                "",
            ]
        )

    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_manifest(chunks: list[ContextChunk], output_path: Path) -> None:
    sources: dict[str, dict[str, Any]] = {}
    for chunk in chunks:
        source = sources.setdefault(
            chunk.source_path,
            {"source_name": chunk.source_name, "source_type": chunk.source_type, "chunks": 0},
        )
        source["chunks"] += 1

    output = {
        "total_sources": len(sources),
        "total_chunks": len(chunks),
        "sources": sources,
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
